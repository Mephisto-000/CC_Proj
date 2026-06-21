"""
Contract Highlighter

對原始合約的副本套用黃色 highlight，標示風險條款。
原始檔案不會被修改。

Usage:
    python highlight_contract.py \
        --input <contract.docx> \
        --findings <findings.json> \
        --output <contract_reviewed.docx>
"""

from __future__ import annotations

import argparse
import json
import logging
import shutil
import sys
from dataclasses import dataclass
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run
except ImportError as exc:
    sys.stderr.write(
        "python-docx 未安裝。請執行：pip install python-docx\n"
    )
    raise SystemExit(1) from exc

logger = logging.getLogger("contract_reviewer.highlight")


@dataclass(frozen=True)
class RiskFinding:
    """從 findings.json 載入的單筆風險紀錄（僅取 highlight 需要的欄位）。"""

    finding_id: str
    para_index: int
    match_text: str
    risk_level: str

    @classmethod
    def from_dict(cls, raw: dict[str, object]) -> "RiskFinding":
        return cls(
            finding_id=str(raw.get("id", "")),
            para_index=int(raw["para_index"]),  # type: ignore[arg-type]
            match_text=str(raw["match_text"]),
            risk_level=str(raw.get("risk_level", "")),
        )


@dataclass(frozen=True)
class HighlightResult:
    finding_id: str
    para_index: int
    success: bool
    detail: str


def _configure_logging(verbose: bool) -> None:
    level = logging.DEBUG if verbose else logging.INFO
    logging.basicConfig(
        level=level,
        format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    )


def _load_findings(findings_path: Path) -> list[RiskFinding]:
    payload = json.loads(findings_path.read_text(encoding="utf-8"))
    risks_raw = payload.get("risks", [])
    if not isinstance(risks_raw, list):
        raise ValueError("findings.json 的 risks 必須為 array")
    return [RiskFinding.from_dict(item) for item in risks_raw]


def _copy_original(input_path: Path, output_path: Path) -> None:
    """先以二進位複製原始檔，再對副本進行修改。確保原檔不被觸碰。"""
    if input_path.resolve() == output_path.resolve():
        raise ValueError("輸出路徑不可與原始檔相同，禁止覆寫原始合約")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(input_path, output_path)
    logger.debug("複製原始檔 %s -> %s", input_path, output_path)


def _clone_run_format(source: Run, target: Run) -> None:
    """把來源 run 的字型格式複製到目標 run。"""
    target.bold = source.bold
    target.italic = source.italic
    target.underline = source.underline
    target.font.name = source.font.name
    target.font.size = source.font.size
    if source.font.color and source.font.color.rgb is not None:
        target.font.color.rgb = source.font.color.rgb


def _split_run_with_highlight(
    paragraph: Paragraph, run_index: int, start: int, end: int
) -> bool:
    """
    將指定 run 切成最多三段：before / matched(highlighted) / after。
    回傳是否成功。
    """
    runs = paragraph.runs
    if run_index >= len(runs):
        return False
    original_run = runs[run_index]
    text = original_run.text
    if start < 0 or end > len(text) or start >= end:
        return False

    before_text = text[:start]
    matched_text = text[start:end]
    after_text = text[end:]

    # 改寫原 run 為 before 段（保留所有格式）
    original_run.text = before_text

    # 在原 run 之後插入 matched run（黃色 highlight）
    element = original_run._element
    matched_run_xml = original_run._element.makeelement(original_run._element.tag, {})

    # 直接用 add_run 新增一個 run 再搬到正確位置
    matched_run = paragraph.add_run(matched_text)
    _clone_run_format(original_run, matched_run)
    matched_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # 把 matched_run 的 XML 從末尾搬到 original_run 之後
    matched_element = matched_run._element
    matched_element.getparent().remove(matched_element)
    element.addnext(matched_element)

    if after_text:
        after_run = paragraph.add_run(after_text)
        _clone_run_format(original_run, after_run)
        after_element = after_run._element
        after_element.getparent().remove(after_element)
        matched_element.addnext(after_element)

    return True


def _highlight_in_paragraph(paragraph: Paragraph, match_text: str) -> bool:
    """
    嘗試在段落中標示 match_text。
    優先嘗試「單一 run 內完整出現」；若失敗則回退到「跨 run 模式」。
    """
    if not match_text:
        return False

    # 1) 單一 run 內完整命中
    for run_index, run in enumerate(paragraph.runs):
        text = run.text
        if not text:
            continue
        position = text.find(match_text)
        if position >= 0:
            return _split_run_with_highlight(
                paragraph, run_index, position, position + len(match_text)
            )

    # 2) 跨 run：把整段重組後比對位置，再回推到各 run
    return _highlight_across_runs(paragraph, match_text)


def _highlight_across_runs(paragraph: Paragraph, match_text: str) -> bool:
    """跨 run 比對：以段落整體文字比對，命中後對涉及的所有 runs 套用 highlight。"""
    runs = paragraph.runs
    if not runs:
        return False

    full_text = "".join(run.text for run in runs)
    position = full_text.find(match_text)
    if position < 0:
        return False

    match_end = position + len(match_text)

    # 計算每個 run 在 full_text 中的起訖
    cursor = 0
    affected: list[tuple[int, int, int]] = []  # (run_index, local_start, local_end)
    for run_index, run in enumerate(runs):
        run_start = cursor
        run_end = cursor + len(run.text)
        if run_end <= position or run_start >= match_end:
            cursor = run_end
            continue
        local_start = max(position, run_start) - run_start
        local_end = min(match_end, run_end) - run_start
        affected.append((run_index, local_start, local_end))
        cursor = run_end

    if not affected:
        return False

    # 簡化策略：對涉及的每個 run，把命中區段套上 highlight。
    # 由於 split 會插入新 run 改變索引，從後往前處理。
    for run_index, local_start, local_end in reversed(affected):
        _split_run_with_highlight(paragraph, run_index, local_start, local_end)
    return True


def apply_highlights(
    target_doc_path: Path, findings: list[RiskFinding]
) -> list[HighlightResult]:
    document = Document(str(target_doc_path))
    paragraphs = document.paragraphs
    results: list[HighlightResult] = []

    for finding in findings:
        if finding.para_index < 0 or finding.para_index >= len(paragraphs):
            results.append(
                HighlightResult(
                    finding_id=finding.finding_id,
                    para_index=finding.para_index,
                    success=False,
                    detail=f"段落索引 {finding.para_index} 超出範圍（共 {len(paragraphs)} 段）",
                )
            )
            continue

        paragraph = paragraphs[finding.para_index]
        ok = _highlight_in_paragraph(paragraph, finding.match_text)

        if not ok:
            # fallback: 全段掃描，找包含 match_text 的段落
            fallback_index = _search_paragraphs_for_match(
                paragraphs, finding.match_text
            )
            if fallback_index is not None:
                ok = _highlight_in_paragraph(
                    paragraphs[fallback_index], finding.match_text
                )
                if ok:
                    results.append(
                        HighlightResult(
                            finding_id=finding.finding_id,
                            para_index=fallback_index,
                            success=True,
                            detail=f"以全文搜尋 fallback 命中（原索引 {finding.para_index}）",
                        )
                    )
                    continue

        results.append(
            HighlightResult(
                finding_id=finding.finding_id,
                para_index=finding.para_index,
                success=ok,
                detail="OK" if ok else "match_text 未在段落中找到",
            )
        )

    document.save(str(target_doc_path))
    return results


def _search_paragraphs_for_match(
    paragraphs: list[Paragraph], match_text: str
) -> int | None:
    for index, paragraph in enumerate(paragraphs):
        if match_text and match_text in paragraph.text:
            return index
        # 跨 run 文字
        if match_text and match_text in "".join(r.text for r in paragraph.runs):
            return index
    return None


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Apply yellow highlights to contract risk clauses.")
    parser.add_argument("--input", type=Path, required=True, help="原始合約 .docx")
    parser.add_argument("--findings", type=Path, required=True, help="findings.json 路徑")
    parser.add_argument("--output", type=Path, required=True, help="輸出的 reviewed.docx")
    parser.add_argument("--verbose", action="store_true", help="顯示 debug log")
    return parser


def main(argv: list[str] | None = None) -> int:
    parser = _build_parser()
    args = parser.parse_args(argv)
    _configure_logging(args.verbose)

    if not args.input.exists():
        logger.error("找不到原始合約：%s", args.input)
        return 2
    if not args.findings.exists():
        logger.error("找不到 findings 檔：%s", args.findings)
        return 2

    try:
        findings = _load_findings(args.findings)
    except (ValueError, KeyError) as exc:
        logger.error("解析 findings 失敗：%s", exc)
        return 2

    try:
        _copy_original(args.input, args.output)
    except ValueError as exc:
        logger.error("%s", exc)
        return 2

    results = apply_highlights(args.output, findings)

    success_count = sum(1 for r in results if r.success)
    fail_count = len(results) - success_count
    logger.info("Highlight 完成：成功 %d / 失敗 %d", success_count, fail_count)

    for result in results:
        if not result.success:
            logger.warning(
                "[%s] para_index=%d 失敗：%s",
                result.finding_id,
                result.para_index,
                result.detail,
            )

    summary = {
        "output": str(args.output),
        "total": len(results),
        "success": success_count,
        "failed": fail_count,
        "results": [
            {
                "id": r.finding_id,
                "para_index": r.para_index,
                "success": r.success,
                "detail": r.detail,
            }
            for r in results
        ],
    }
    sys.stdout.write(json.dumps(summary, ensure_ascii=False, indent=2) + "\n")
    return 0 if fail_count == 0 else 3


if __name__ == "__main__":
    raise SystemExit(main())
