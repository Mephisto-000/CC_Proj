"""
Contract Extractor

掃描 .docx 合約或目錄，輸出結構化 JSON 供後續審查使用。

Usage:
    python extract_contract.py --scan <dir>
    python extract_contract.py --input <contract.docx> --output <extracted.json>
"""

from __future__ import annotations

import argparse
import json
import logging
import sys
from dataclasses import asdict, dataclass, field
from pathlib import Path

try:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError as exc:
    sys.stderr.write(
        "python-docx 未安裝。請執行：pip install python-docx\n"
    )
    raise SystemExit(1) from exc

logger = logging.getLogger("contract_reviewer.extract")

SUPPORTED_EXTENSIONS: tuple[str, ...] = (".docx", ".doc")


@dataclass(frozen=True)
class ParagraphRecord:
    para_index: int
    text: str
    style: str
    is_heading: bool


@dataclass(frozen=True)
class TableCellRecord:
    row: int
    col: int
    text: str


@dataclass(frozen=True)
class TableRecord:
    table_index: int
    cells: list[TableCellRecord]


@dataclass(frozen=True)
class ContractMetadata:
    file_name: str
    paragraph_count: int
    table_count: int


@dataclass
class ExtractedContract:
    metadata: ContractMetadata
    paragraphs: list[ParagraphRecord] = field(default_factory=list)
    tables: list[TableRecord] = field(default_factory=list)

    def to_dict(self) -> dict[str, object]:
        return {
            "metadata": asdict(self.metadata),
            "paragraphs": [asdict(p) for p in self.paragraphs],
            "tables": [
                {
                    "table_index": t.table_index,
                    "cells": [asdict(c) for c in t.cells],
                }
                for t in self.tables
            ],
        }


def _configure_logging(verbose: bool) -> None:
    level = logging.DEBUG if verbose else logging.INFO
    logging.basicConfig(
        level=level,
        format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    )


def scan_directory(target_dir: Path) -> list[Path]:
    """回傳資料夾中所有支援的合約檔案路徑（不含隱藏檔與 reviewed/risk_report 衍生檔）。"""
    if not target_dir.is_dir():
        raise NotADirectoryError(f"非目錄：{target_dir}")

    results: list[Path] = []
    for path in sorted(target_dir.iterdir()):
        if not path.is_file():
            continue
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        if path.name.startswith("~$") or path.name.startswith("."):
            continue
        stem = path.stem
        if stem.endswith("_reviewed") or stem.endswith("_risk_report"):
            continue
        results.append(path)
    return results


def _is_heading_style(style_name: str) -> bool:
    if not style_name:
        return False
    lowered = style_name.lower()
    return "heading" in lowered or "標題" in style_name


def _extract_paragraphs(document: DocumentType) -> list[ParagraphRecord]:
    records: list[ParagraphRecord] = []
    for index, paragraph in enumerate(document.paragraphs):
        paragraph_obj: Paragraph = paragraph
        text = paragraph_obj.text
        style_name = paragraph_obj.style.name if paragraph_obj.style else ""
        records.append(
            ParagraphRecord(
                para_index=index,
                text=text,
                style=style_name,
                is_heading=_is_heading_style(style_name),
            )
        )
    return records


def _extract_tables(document: DocumentType) -> list[TableRecord]:
    tables: list[TableRecord] = []
    for table_index, table in enumerate(document.tables):
        table_obj: Table = table
        cells: list[TableCellRecord] = []
        for row_index, row in enumerate(table_obj.rows):
            for col_index, cell in enumerate(row.cells):
                cells.append(
                    TableCellRecord(
                        row=row_index,
                        col=col_index,
                        text=cell.text,
                    )
                )
        tables.append(TableRecord(table_index=table_index, cells=cells))
    return tables


def extract_contract(input_path: Path) -> ExtractedContract:
    if input_path.suffix.lower() == ".doc":
        raise ValueError(
            f".doc 格式需先用 LibreOffice 轉成 .docx：{input_path.name}"
        )
    if not input_path.exists():
        raise FileNotFoundError(input_path)

    logger.info("讀取合約：%s", input_path.name)
    document = Document(str(input_path))
    paragraphs = _extract_paragraphs(document)
    tables = _extract_tables(document)
    metadata = ContractMetadata(
        file_name=input_path.name,
        paragraph_count=len(paragraphs),
        table_count=len(tables),
    )
    return ExtractedContract(
        metadata=metadata, paragraphs=paragraphs, tables=tables
    )


def _write_json(payload: dict[str, object], output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    logger.info("已寫出：%s", output_path)


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Extract contract content for reviewing.")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--scan", type=Path, help="掃描目錄並列出合約檔")
    group.add_argument("--input", type=Path, help="單一合約輸入路徑")
    parser.add_argument("--output", type=Path, help="輸出 JSON 路徑（--input 模式必填）")
    parser.add_argument("--verbose", action="store_true", help="顯示 debug log")
    return parser


def main(argv: list[str] | None = None) -> int:
    parser = _build_parser()
    args = parser.parse_args(argv)
    _configure_logging(args.verbose)

    if args.scan is not None:
        try:
            files = scan_directory(args.scan)
        except NotADirectoryError as exc:
            logger.error("%s", exc)
            return 2
        result = {
            "directory": str(args.scan),
            "count": len(files),
            "files": [f.name for f in files],
        }
        sys.stdout.write(json.dumps(result, ensure_ascii=False, indent=2) + "\n")
        return 0

    if args.input is None:
        logger.error("--input 為必填")
        return 2
    if args.output is None:
        logger.error("--input 模式必須提供 --output")
        return 2

    try:
        extracted = extract_contract(args.input)
    except (FileNotFoundError, ValueError) as exc:
        logger.error("%s", exc)
        return 2

    _write_json(extracted.to_dict(), args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
