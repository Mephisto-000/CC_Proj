"""
Risk Report Generator

依 findings.json 產生 Word 風險分析報告。

Usage:
    python generate_risk_report.py \
        --findings <findings.json> \
        --contract-name <原始檔名> \
        --output <contract_risk_report.docx>
"""

from __future__ import annotations

import argparse
import json
import logging
import sys
from collections import Counter
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor
except ImportError as exc:
    sys.stderr.write(
        "python-docx 未安裝。請執行：pip install python-docx\n"
    )
    raise SystemExit(1) from exc

logger = logging.getLogger("contract_reviewer.report")


CATEGORY_LABELS: dict[str, str] = {
    "unreasonable_liability": "不合理責任",
    "unilateral_termination": "單方終止權",
    "payment_risk": "付款風險",
    "acceptance_risk": "驗收風險",
    "ip_risk": "智慧財產權風險",
    "confidentiality_risk": "保密義務風險",
    "penalty_risk": "違約金風險",
}

RISK_LEVEL_COLORS: dict[str, RGBColor] = {
    "高": RGBColor(0xC0, 0x39, 0x2B),  # 深紅
    "中": RGBColor(0xD3, 0x88, 0x00),  # 橘
    "低": RGBColor(0x16, 0xA0, 0x85),  # 綠
}

RISK_LEVEL_ORDER: dict[str, int] = {"高": 0, "中": 1, "低": 2}


@dataclass(frozen=True)
class RiskItem:
    finding_id: str
    clause_id: str
    para_index: int
    match_text: str
    category: str
    risk_level: str
    reason: str
    suggestion: str

    @property
    def category_label(self) -> str:
        return CATEGORY_LABELS.get(self.category, self.category)

    @classmethod
    def from_dict(cls, raw: dict[str, object]) -> "RiskItem":
        return cls(
            finding_id=str(raw.get("id", "")),
            clause_id=str(raw.get("clause_id", "")),
            para_index=int(raw.get("para_index", -1)),  # type: ignore[arg-type]
            match_text=str(raw.get("match_text", "")),
            category=str(raw.get("category", "")),
            risk_level=str(raw.get("risk_level", "")),
            reason=str(raw.get("reason", "")),
            suggestion=str(raw.get("suggestion", "")),
        )


@dataclass(frozen=True)
class MissingClause:
    clause_id: str
    name: str
    reason: str
    suggestion: str

    @classmethod
    def from_dict(cls, raw: dict[str, object]) -> "MissingClause":
        return cls(
            clause_id=str(raw.get("id", "")),
            name=str(raw.get("name", "")),
            reason=str(raw.get("reason", "")),
            suggestion=str(raw.get("suggestion", "")),
        )


@dataclass
class Findings:
    contract_name: str
    review_date: str
    party_a_name: str
    party_b_name: str
    risks: list[RiskItem] = field(default_factory=list)
    missing_clauses: list[MissingClause] = field(default_factory=list)

    @classmethod
    def from_dict(cls, raw: dict[str, object]) -> "Findings":
        risks_raw = raw.get("risks", [])
        missing_raw = raw.get("missing_clauses", [])
        if not isinstance(risks_raw, list) or not isinstance(missing_raw, list):
            raise ValueError("risks / missing_clauses 必須為 array")
        return cls(
            contract_name=str(raw.get("contract_name", "")),
            review_date=str(raw.get("review_date", date.today().isoformat())),
            party_a_name=str(raw.get("party_a_name", "")),
            party_b_name=str(raw.get("party_b_name", "")),
            risks=[RiskItem.from_dict(r) for r in risks_raw],
            missing_clauses=[MissingClause.from_dict(m) for m in missing_raw],
        )


def _configure_logging(verbose: bool) -> None:
    level = logging.DEBUG if verbose else logging.INFO
    logging.basicConfig(
        level=level,
        format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    )


def _add_heading(document: object, text: str, level: int) -> None:
    heading = document.add_heading(text, level=level)  # type: ignore[attr-defined]
    for run in heading.runs:
        run.font.name = "Microsoft JhengHei"


def _add_paragraph(document: object, text: str, bold: bool = False) -> None:
    paragraph = document.add_paragraph()  # type: ignore[attr-defined]
    run = paragraph.add_run(text)
    run.font.name = "Microsoft JhengHei"
    run.font.size = Pt(11)
    run.bold = bold


def _set_cell_text(cell: object, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""  # type: ignore[attr-defined]
    paragraph = cell.paragraphs[0]  # type: ignore[attr-defined]
    run = paragraph.add_run(text)
    run.font.name = "Microsoft JhengHei"
    run.font.size = Pt(10)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP  # type: ignore[attr-defined]


def _build_header(document: object, findings: Findings) -> None:
    _add_heading(document, "合約風險分析報告", level=0)

    info_table = document.add_table(rows=4, cols=2)  # type: ignore[attr-defined]
    info_table.style = "Light Grid Accent 1"
    info_table.autofit = False
    info_table.columns[0].width = Cm(4)
    info_table.columns[1].width = Cm(12)

    info_rows: tuple[tuple[str, str], ...] = (
        ("合約名稱", findings.contract_name),
        ("審查日期", findings.review_date),
        ("甲方", findings.party_a_name),
        ("乙方（本次審查立場）", findings.party_b_name),
    )
    for row_index, (label, value) in enumerate(info_rows):
        _set_cell_text(info_table.rows[row_index].cells[0], label, bold=True)
        _set_cell_text(info_table.rows[row_index].cells[1], value)

    document.add_paragraph()  # type: ignore[attr-defined]


def _build_summary(document: object, findings: Findings) -> None:
    _add_heading(document, "一、風險摘要", level=1)

    level_counts = Counter(r.risk_level for r in findings.risks)
    category_counts = Counter(r.category for r in findings.risks)

    _add_paragraph(
        document,
        f"本次共識別 {len(findings.risks)} 筆風險條款，"
        f"另發現 {len(findings.missing_clauses)} 項缺失之保護性條款。",
    )

    # 風險等級分布
    summary_table = document.add_table(rows=2, cols=4)  # type: ignore[attr-defined]
    summary_table.style = "Light Grid Accent 1"
    headers = ("高", "中", "低", "合計")
    for col_index, header in enumerate(headers):
        _set_cell_text(summary_table.rows[0].cells[col_index], header, bold=True)
    values = (
        str(level_counts.get("高", 0)),
        str(level_counts.get("中", 0)),
        str(level_counts.get("低", 0)),
        str(len(findings.risks)),
    )
    for col_index, value in enumerate(values):
        _set_cell_text(summary_table.rows[1].cells[col_index], value)

    document.add_paragraph()  # type: ignore[attr-defined]

    if category_counts:
        _add_paragraph(document, "依風險類別分布：", bold=True)
        cat_table = document.add_table(rows=len(category_counts) + 1, cols=2)  # type: ignore[attr-defined]
        cat_table.style = "Light Grid Accent 1"
        _set_cell_text(cat_table.rows[0].cells[0], "類別", bold=True)
        _set_cell_text(cat_table.rows[0].cells[1], "筆數", bold=True)
        for row_index, (category, count) in enumerate(
            sorted(category_counts.items(), key=lambda kv: -kv[1]), start=1
        ):
            label = CATEGORY_LABELS.get(category, category)
            _set_cell_text(cat_table.rows[row_index].cells[0], label)
            _set_cell_text(cat_table.rows[row_index].cells[1], str(count))

    document.add_paragraph()  # type: ignore[attr-defined]


def _build_risk_details(document: object, findings: Findings) -> None:
    _add_heading(document, "二、風險條款明細", level=1)

    if not findings.risks:
        _add_paragraph(document, "本次審查未發現需標示之風險條款。")
        return

    sorted_risks = sorted(
        findings.risks,
        key=lambda r: (
            RISK_LEVEL_ORDER.get(r.risk_level, 99),
            r.finding_id,
        ),
    )

    for risk in sorted_risks:
        _add_heading(document, f"{risk.finding_id} ｜ {risk.clause_id}", level=2)

        table = document.add_table(rows=5, cols=2)  # type: ignore[attr-defined]
        table.style = "Light Grid Accent 1"
        table.autofit = False
        table.columns[0].width = Cm(3.5)
        table.columns[1].width = Cm(13)

        rows_def: tuple[tuple[str, str, RGBColor | None], ...] = (
            ("風險等級", risk.risk_level, RISK_LEVEL_COLORS.get(risk.risk_level)),
            ("風險類別", risk.category_label, None),
            ("條款原文", risk.match_text, None),
            ("不利原因", risk.reason, None),
            ("修改建議", risk.suggestion, None),
        )
        for row_index, (label, value, color) in enumerate(rows_def):
            _set_cell_text(table.rows[row_index].cells[0], label, bold=True)
            _set_cell_text(table.rows[row_index].cells[1], value, color=color)

        document.add_paragraph()  # type: ignore[attr-defined]


def _build_missing_clauses(document: object, findings: Findings) -> None:
    _add_heading(document, "三、缺失之保護性條款", level=1)

    if not findings.missing_clauses:
        _add_paragraph(document, "本合約已涵蓋常見保護性條款，無重大缺漏。")
        return

    for clause in findings.missing_clauses:
        _add_heading(document, f"{clause.clause_id} ｜ {clause.name}", level=2)
        table = document.add_table(rows=2, cols=2)  # type: ignore[attr-defined]
        table.style = "Light Grid Accent 1"
        table.autofit = False
        table.columns[0].width = Cm(3.5)
        table.columns[1].width = Cm(13)
        rows_def: tuple[tuple[str, str], ...] = (
            ("缺失影響", clause.reason),
            ("建議補充", clause.suggestion),
        )
        for row_index, (label, value) in enumerate(rows_def):
            _set_cell_text(table.rows[row_index].cells[0], label, bold=True)
            _set_cell_text(table.rows[row_index].cells[1], value)
        document.add_paragraph()  # type: ignore[attr-defined]


def _build_recommendations(document: object, findings: Findings) -> None:
    _add_heading(document, "四、整體建議", level=1)

    high_count = sum(1 for r in findings.risks if r.risk_level == "高")
    medium_count = sum(1 for r in findings.risks if r.risk_level == "中")

    if high_count > 0:
        _add_paragraph(
            document,
            f"本合約存在 {high_count} 筆高風險條款，建議於簽約前完成修改，"
            "否則乙方可能承擔不對等的法律與財務風險。",
        )
    if medium_count > 0:
        _add_paragraph(
            document,
            f"另有 {medium_count} 筆中度風險條款，建議於談判時提出修正，"
            "若甲方堅持原條文，需綜合評估專案商業價值再決定是否承擔。",
        )
    if findings.missing_clauses:
        _add_paragraph(
            document,
            f"缺失 {len(findings.missing_clauses)} 項保護性條款，"
            "應於合約增訂相關條款，避免日後爭議。",
        )
    if high_count == 0 and medium_count == 0 and not findings.missing_clauses:
        _add_paragraph(
            document,
            "整體條款結構對乙方尚屬合理，僅有少量建議事項，可進入簽約程序。",
        )

    _add_paragraph(document, "本報告由 contract-reviewer skill 自動產生，僅供決策參考；正式簽約前建議由法務人員複核。")


def generate_report(findings: Findings, output_path: Path) -> None:
    document = Document()

    # 設定預設字型
    style = document.styles["Normal"]
    style.font.name = "Microsoft JhengHei"
    style.font.size = Pt(11)

    _build_header(document, findings)
    _build_summary(document, findings)
    _build_risk_details(document, findings)
    _build_missing_clauses(document, findings)
    _build_recommendations(document, findings)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    logger.info("風險報告已輸出：%s", output_path)


def _load_findings(findings_path: Path) -> Findings:
    raw = json.loads(findings_path.read_text(encoding="utf-8"))
    return Findings.from_dict(raw)


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Generate contract risk analysis report.")
    parser.add_argument("--findings", type=Path, required=True, help="findings.json 路徑")
    parser.add_argument("--contract-name", type=str, help="覆寫合約名稱（可選）")
    parser.add_argument("--output", type=Path, required=True, help="輸出 .docx 路徑")
    parser.add_argument("--verbose", action="store_true", help="顯示 debug log")
    return parser


def main(argv: list[str] | None = None) -> int:
    parser = _build_parser()
    args = parser.parse_args(argv)
    _configure_logging(args.verbose)

    if not args.findings.exists():
        logger.error("找不到 findings 檔：%s", args.findings)
        return 2

    try:
        findings = _load_findings(args.findings)
    except (ValueError, KeyError) as exc:
        logger.error("解析 findings 失敗：%s", exc)
        return 2

    if args.contract_name:
        findings.contract_name = args.contract_name

    generate_report(findings, args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
